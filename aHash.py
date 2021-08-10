#inbuilt average hash

#extracting images from the corresponding folders
import cv2
import os
from PIL import Image
import imagehash
import matplotlib.pyplot as plt
import xlsxwriter 
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import math

def load_images_from_folder(folder):
	ImgList = [];
	for file in os.listdir(folder):
		if file.endswith('.jpeg') or file.endswith('.jpg') or file.endswith('.png'):
			ImgList.append(file);
	return ImgList;

folder1 = "C:\\Users\\Dell\\OneDrive\\Desktop\\PROJECT\\UserData";
folder2 = "C:\\Users\\Dell\\OneDrive\\Desktop\\PROJECT\\Database";
ImgList1 = load_images_from_folder(folder1);
ImgList2 = load_images_from_folder(folder2);
#print(ImgList1);
#print(ImgList2);

#saving the data to excel sheet
workbook = xlsxwriter.Workbook('C:\\Users\\Dell\\OneDrive\\Desktop\\PROJECT\\RESULT\\Result.xlsx');
DataList = {};
AnswerList = [];
AnswerList1 = [];
for img1 in ImgList1: 
	hash1 = imagehash.average_hash(Image.open('UserData/' + img1));
	d2 = {};
	worksheet = workbook.add_worksheet(img1) 
	row = 0
	col = 0
	worksheet.write(row, col, "Image")
	worksheet.write(row, col + 1, "Hash-diff")  
	row += 1
	for img2 in ImgList2:
		hash2 = imagehash.average_hash(Image.open('Database/' + img2));	
		d2[img2] = hash1 - hash2;
		if hash1 - hash2 == 0:
			AnswerList.append(img1);
			AnswerList1.append(img2);
	DataList[img1] = d2;
	#print(d2);
	#plot a graph
	x = d2.keys();
	y = d2.values();	
	# plotting the grpahs
	low = min(y)
	high = max(y)
	plt.ylim([math.ceil(low-0.5*(high-low)), math.ceil(high+0.5*(high-low))])
	rects = plt.bar(x , y , color='b');
	plt.xlabel("DB Images")
	plt.ylabel("Hash - Difference")
	plt.title('Hash Value Difference-Graph for ' + str(img1))
	for rect in rects:
        	height = rect.get_height()
        	plt.text(rect.get_x() + rect.get_width()/2., 1.05*height,'%d' % int(height),ha='center', va='bottom')
	plt.show();
	for key, value in d2.items():
		worksheet.write(row, col, key) 
		worksheet.write(row, col + 1, value) 
		row += 1
	
#print(DataList);
#print(HashList);
#print("AnswerList");
#print(AnswerList);
ans = [];
for img in ImgList1:
	ans.append(img);	
plt.legend(ans , loc = 4)
#plt.show();
workbook.close();

document = Document()
document.add_heading('THE FOLLOWING IMAGES FROM THE SUBMITTED ASSIGNMENT ARE MATCHED WITH THE IMAGES IN THE DATABASE' , 0)  
document.add_heading('PLAGIARISED IMAGES || DATABASE IMAGES', 4)
#doc = docx.Document("C:\\Users\\Dell\\OneDrive\\Desktop\\PROJECT\\RESULT\\PlagiarisedImages.docx") 
index = 0;
paragraph = document.add_paragraph();
for img in AnswerList :
	#im = Image.open('UserData/' + img)  
	#im.show()
	#paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
	run = paragraph.add_run()
	run.add_break()
	run.add_picture('UserData/' + img , width = Inches(1.25) , height = Inches(1.25))
	#paragraph_format.space_before = Pt(18)
	#paragraph_format.tab_stops
	run_2 = paragraph.add_run()
	run_2.add_picture('Database/' + AnswerList1[index] , width = Inches(1.0) , height = Inches(1.0)) 
	index = index + 1;
document.save('C:\\Users\\Dell\\OneDrive\\Desktop\\PROJECT\\RESULT\\PlagiarisedImages.docx')
